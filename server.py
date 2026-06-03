from __future__ import annotations

import mimetypes
import os
import re
import shutil
import socket
import tempfile
from html import unescape
from pathlib import Path
from typing import Any
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qs, urlparse
from urllib.request import Request, urlopen

from flask import Flask, after_this_request, jsonify, request, send_file, send_from_directory
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from pypdf import PdfReader

WORKSPACE_DIR = Path(__file__).resolve().parent
PUBLIC_DIR = WORKSPACE_DIR / "public"
DOWNLOADS_DIR = WORKSPACE_DIR / "downloads"
STATIC_FILES = {
    "index.html": "text/html; charset=utf-8",
    "styles.css": "text/css; charset=utf-8",
    "script.js": "application/javascript; charset=utf-8",
}

USER_AGENT = (
    "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
    "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/124.0 Safari/537.36"
)

app = Flask(__name__, static_folder=None)


class DownloadError(RuntimeError):
    pass


def resolve_static_file(file_name: str) -> Path:
    public_path = PUBLIC_DIR / file_name
    if public_path.exists():
        return public_path

    root_path = WORKSPACE_DIR / file_name
    if root_path.exists():
        return root_path

    raise FileNotFoundError(file_name)


def sanitize_file_name(value: str, fallback: str = "downloaded-file") -> str:
    cleaned = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", value or "").strip().strip(".")
    return cleaned or fallback


def send_request(url: str) -> tuple[dict[str, str], bytes]:
    request_obj = Request(url, headers={"User-Agent": USER_AGENT})
    try:
        with urlopen(request_obj, timeout=90) as response:
            headers = {key.lower(): value for key, value in response.info().items()}
            return headers, response.read()
    except HTTPError as exc:
        raise DownloadError(f"目标链接返回了 {exc.code}。") from exc
    except URLError as exc:
        raise DownloadError("无法连接到目标站点，请稍后再试。") from exc


def strip_tags(value: str) -> str:
    return re.sub(r"<[^>]+>", "", value)


def parse_yongzin_title(html_text: str, fallback: str) -> str:
    heading_match = re.search(r"<h3[^>]*>(.*?)</h3>", html_text, re.IGNORECASE | re.DOTALL)
    if heading_match:
        title = sanitize_file_name(unescape(strip_tags(heading_match.group(1))).strip(), fallback)
        if title:
            return title

    title_match = re.search(r"<title>(.*?)</title>", html_text, re.IGNORECASE | re.DOTALL)
    if title_match:
        title = sanitize_file_name(unescape(strip_tags(title_match.group(1))).strip(), fallback)
        if title:
            return title

    return fallback


def normalize_blocks(text: str) -> list[str]:
    lines = [line.rstrip() for line in text.replace("\r", "").split("\n")]
    blocks: list[str] = []
    current: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if current:
                blocks.append("\n".join(current))
                current = []
            continue
        current.append(stripped)

    if current:
        blocks.append("\n".join(current))

    return blocks


def apply_style_font(style: Any, font_name: str, font_size: int = 14) -> None:
    style.font.name = font_name
    style.font.size = Pt(font_size)

    style_props = style._element.get_or_add_rPr()
    fonts = style_props.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        style_props.append(fonts)

    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)
    fonts.set(qn("w:cs"), font_name)

    lang = style_props.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        style_props.append(lang)

    lang.set(qn("w:val"), "bo-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    lang.set(qn("w:bidi"), "bo-CN")


def apply_run_font(run: Any, font_name: str, font_size: int = 14) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)

    run_props = run._element.get_or_add_rPr()
    fonts = run_props.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        run_props.append(fonts)

    fonts.set(qn("w:ascii"), font_name)
    fonts.set(qn("w:hAnsi"), font_name)
    fonts.set(qn("w:eastAsia"), font_name)
    fonts.set(qn("w:cs"), font_name)

    lang = run_props.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        run_props.append(lang)

    lang.set(qn("w:val"), "bo-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")
    lang.set(qn("w:bidi"), "bo-CN")


def extract_pdf_pages(pdf_path: Path) -> list[list[str]]:
    reader = PdfReader(str(pdf_path))
    pages: list[list[str]] = []

    for page in reader.pages:
        text = page.extract_text() or ""
        blocks = normalize_blocks(text)
        if blocks:
            pages.append(blocks)

    return pages


def create_docx_from_pages(
    pages: list[list[str]],
    destination: Path,
    title: str,
    source_url: str,
    source_pdf_name: str,
) -> None:
    font_name = "Microsoft Himalaya"
    document = Document()
    apply_style_font(document.styles["Normal"], font_name, 14)

    heading = document.add_paragraph()
    heading_run = heading.add_run(title)
    heading_run.bold = True
    apply_run_font(heading_run, font_name, 18)
    heading.paragraph_format.space_after = Pt(8)

    intro = document.add_paragraph()
    intro_run = intro.add_run(f"来源链接：{source_url}\n备份 PDF：{source_pdf_name}")
    apply_run_font(intro_run, font_name, 10)
    intro.paragraph_format.space_after = Pt(12)

    first_page = True
    for page_blocks in pages:
        if not first_page:
            document.add_page_break()
        first_page = False

        for block in page_blocks:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(block)
            apply_run_font(run, font_name, 14)

    document.save(destination)


def maybe_cache_artifact(artifact_path: Path, download_name: str) -> None:
    try:
        DOWNLOADS_DIR.mkdir(exist_ok=True)
        shutil.copyfile(artifact_path, DOWNLOADS_DIR / download_name)
    except OSError:
        pass


def build_result(path: Path, download_name: str, mime_type: str, message: str, temp_dir: Path) -> dict[str, Any]:
    maybe_cache_artifact(path, download_name)
    return {
        "path": path,
        "download_name": download_name,
        "mime_type": mime_type,
        "message": message,
        "temp_dir": temp_dir,
    }


def prepare_yongzin_download(url: str, requested_name: str, requested_format: str, temp_dir: Path) -> dict[str, Any]:
    parsed = urlparse(url)
    query = parse_qs(parsed.query)
    doc_id = (query.get("id") or query.get("docNo") or [""])[0].strip()
    if not doc_id:
        raise DownloadError("没有在链接里找到文档编号。")

    base_name = sanitize_file_name(Path(requested_name).stem if requested_name else doc_id, doc_id)
    html_headers, html_body = send_request(url)
    del html_headers
    html_text = html_body.decode("utf-8", errors="ignore")
    title = parse_yongzin_title(html_text, base_name)
    pdf_url = f"https://wenku.yongzin.com/topLayoutController/getPdfFile.do?wenku_uuid={doc_id}"
    pdf_headers, pdf_body = send_request(pdf_url)

    pdf_name = sanitize_file_name(f"{title}.pdf", f"{base_name}.pdf")
    pdf_path = temp_dir / pdf_name
    pdf_path.write_bytes(pdf_body)

    if requested_format == "pdf":
        return build_result(
            pdf_path,
            pdf_name,
            pdf_headers.get("content-type", "application/pdf"),
            "已生成 PDF 下载。",
            temp_dir,
        )

    pages = extract_pdf_pages(pdf_path)
    combined_text = "".join(block for page in pages for block in page)
    if len(combined_text.strip()) < 20:
        return build_result(
            pdf_path,
            pdf_name,
            "application/pdf",
            "文档文字层不足，已自动退回为 PDF 下载。",
            temp_dir,
        )

    docx_name = sanitize_file_name(f"{title}.docx", f"{base_name}.docx")
    docx_path = temp_dir / docx_name
    create_docx_from_pages(pages, docx_path, title, url, pdf_name)

    return build_result(
        docx_path,
        docx_name,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "已生成更适合 WPS 打开的 Word 文本版。",
        temp_dir,
    )


def prepare_direct_download(url: str, requested_name: str, temp_dir: Path) -> dict[str, Any]:
    parsed = urlparse(url)
    guessed_name = Path(parsed.path).name or "downloaded-file"
    _, body = send_request(url)
    mime_type = mimetypes.guess_type(guessed_name)[0] or "application/octet-stream"
    download_name = sanitize_file_name(requested_name or guessed_name, "downloaded-file")
    destination = temp_dir / download_name
    destination.write_bytes(body)
    return build_result(destination, download_name, mime_type, "文件已处理完成。", temp_dir)


def prepare_download(url: str, requested_name: str, requested_format: str) -> dict[str, Any]:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"}:
        raise DownloadError("目前只支持 http 或 https 链接。")

    temp_dir = Path(tempfile.mkdtemp(prefix="yuncang-"))

    if parsed.netloc == "wenku.yongzin.com" and parsed.path.endswith("/fileInfo.do"):
        effective_format = "docx" if requested_format == "auto" else requested_format
        return prepare_yongzin_download(url, requested_name, effective_format, temp_dir)

    return prepare_direct_download(url, requested_name, temp_dir)


def get_lan_addresses() -> list[str]:
    addresses: set[str] = set()

    try:
        _, _, host_addresses = socket.gethostbyname_ex(socket.gethostname())
        for address in host_addresses:
            if address and not address.startswith("127."):
                addresses.add(address)
    except OSError:
        pass

    try:
        with socket.socket(socket.AF_INET, socket.SOCK_DGRAM) as sock:
            sock.connect(("8.8.8.8", 80))
            address = sock.getsockname()[0]
            if address and not address.startswith("127."):
                addresses.add(address)
    except OSError:
        pass

    return sorted(addresses)


def build_server_info() -> dict[str, Any]:
    origin = request.host_url.rstrip("/")
    host = request.host.split(":", 1)[0]
    port = request.environ.get("SERVER_PORT", "4173")

    if host in {"127.0.0.1", "localhost"}:
        lan_urls = [f"http://{address}:{port}" for address in get_lan_addresses()]
        access_links = [{"label": "本机使用", "url": origin}]
        access_links.extend({"label": "发给同一局域网的人", "url": url} for url in lan_urls)
        return {
            "shareHint": "127.0.0.1 只对你自己有效；如果要同一 Wi-Fi 下共享，用下面的局域网地址。",
            "accessLinks": access_links,
        }

    return {
        "shareHint": "这是当前可公开访问的网址，发给别人后，对方用 Wi-Fi 或流量都能打开。",
        "accessLinks": [{"label": "公开网址", "url": origin}],
    }


@app.get("/")
@app.get("/index.html")
def serve_index():
    file_path = resolve_static_file("index.html")
    return send_file(file_path, mimetype=STATIC_FILES["index.html"])


@app.get("/styles.css")
def serve_styles():
    file_path = resolve_static_file("styles.css")
    return send_file(file_path, mimetype=STATIC_FILES["styles.css"])


@app.get("/script.js")
def serve_script():
    file_path = resolve_static_file("script.js")
    return send_file(file_path, mimetype=STATIC_FILES["script.js"])


@app.get("/api/server-info")
def server_info():
    return jsonify(build_server_info())


@app.post("/api/download")
def download():
    payload = request.get_json(silent=True) or {}
    url = str(payload.get("url", "")).strip()
    requested_name = sanitize_file_name(str(payload.get("fileName", "")).strip(), "")
    requested_format = str(payload.get("format", "auto")).strip() or "auto"

    if not url:
        return jsonify({"success": False, "message": "请先提供一个链接。"}), 400

    try:
        result = prepare_download(url, requested_name, requested_format)
    except DownloadError as exc:
        return jsonify({"success": False, "message": str(exc)}), 400
    except Exception as exc:  # noqa: BLE001
        return jsonify({"success": False, "message": f"处理失败：{exc}"}), 500

    temp_dir = result["temp_dir"]

    @after_this_request
    def cleanup(response):
        shutil.rmtree(temp_dir, ignore_errors=True)
        return response

    response = send_file(
        result["path"],
        as_attachment=True,
        download_name=result["download_name"],
        mimetype=result["mime_type"],
        max_age=0,
    )
    response.headers["X-Yuncang-Message"] = result["message"]
    response.headers["Cache-Control"] = "no-store"
    return response


if __name__ == "__main__":
    app.run(
        host=os.environ.get("YUNCANG_HOST", "127.0.0.1"),
        port=int(os.environ.get("YUNCANG_PORT", "4173")),
        debug=False,
    )
